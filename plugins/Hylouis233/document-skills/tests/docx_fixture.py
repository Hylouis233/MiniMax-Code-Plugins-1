# Minimal runnable fixtures for the snippets called out in review: one file per format.
# Each script is self-contained, writes only scratch files into the current directory,
# and exits non-zero on failed assertions. Run from any scratch directory:
#   python pdf_fixture.py   (deps: reportlab, pypdf, pymupdf)
#   python pptx_fixture.py  (deps: python-pptx)
#   python xlsx_fixture.py  (deps: openpyxl)
#   python docx_fixture.py  (deps: python-docx, pymupdf, soffice on PATH)
import copy
import subprocess
import sys
import zipfile

import fitz
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree

failures = []


def check(name, cond, extra=""):
    print(("PASS " if cond else "FAIL ") + name + ((" :: " + str(extra)) if not cond and extra else ""))
    if not cond:
        failures.append(name)


# ---- review.md health check: external entities stay unresolved ---------------
safe_xml_parser = etree.XMLParser(
    load_dtd=False,
    resolve_entities=False,
    no_network=True,
    huge_tree=False,
    recover=False,
)
hostile_xml = (
    b'<!DOCTYPE root [<!ENTITY xxe SYSTEM "file:///definitely-not-readable">]>'
    b'<root>&xxe;</root>'
)
parsed = etree.fromstring(hostile_xml, parser=safe_xml_parser)
check("DOCX XML parser leaves external entities unresolved", parsed.text is None and len(parsed) == 1)

# ---- review.md health check rejects archive bombs before expanding parts -------
MAX_XML_PART = 20 * 1024 * 1024
MAX_ENTRY = 100 * 1024 * 1024
MAX_TOTAL_UNCOMPRESSED = 500 * 1024 * 1024
MAX_COMPRESSION_RATIO = 200


def validate_docx_package(path):
    with zipfile.ZipFile(path) as archive:
        infos = archive.infolist()
        names = {info.filename for info in infos}
        assert len(names) == len(infos)
        assert "[Content_Types].xml" in names and "word/document.xml" in names
        assert sum(info.file_size for info in infos) <= MAX_TOTAL_UNCOMPRESSED
        actual_total = 0
        for info in infos:
            assert info.file_size <= MAX_ENTRY
            assert info.file_size / max(info.compress_size, 1) <= MAX_COMPRESSION_RATIO
            is_xml = info.filename.endswith((".xml", ".rels"))
            if is_xml:
                assert info.file_size <= MAX_XML_PART
            chunks = []
            actual_size = 0
            with archive.open(info) as stream:
                while chunk := stream.read(64 * 1024):
                    actual_size += len(chunk)
                    actual_total += len(chunk)
                    assert actual_size <= MAX_ENTRY
                    assert actual_total <= MAX_TOTAL_UNCOMPRESSED
                    if is_xml:
                        chunks.append(chunk)
            assert actual_size == info.file_size
            if is_xml:
                etree.fromstring(b"".join(chunks), parser=safe_xml_parser)


health_doc = Document()
health_doc.add_paragraph("bounded health check")
health_doc.save("healthy.docx")
try:
    validate_docx_package("healthy.docx")
    healthy_package_passed = True
except Exception:
    healthy_package_passed = False
check("bounded package health check accepts an ordinary DOCX", healthy_package_passed)

with zipfile.ZipFile("compressed-bomb.docx", "w", zipfile.ZIP_DEFLATED) as archive:
    archive.writestr("[Content_Types].xml", "<Types/>")
    archive.writestr("word/document.xml", "<document>" + (" " * 2_000_000) + "</document>")
try:
    validate_docx_package("compressed-bomb.docx")
    archive_bomb_rejected = False
except AssertionError:
    archive_bomb_rejected = True
check("suspicious compression ratio is rejected before XML expansion", archive_bomb_rejected)

# ---- read.md includes block paragraphs nested in content controls --------------
def iter_body_paragraphs(parent, document):
    for child in parent.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:sdt"):
            content = child.find(qn("w:sdtContent"))
            if content is not None:
                yield from iter_body_paragraphs(content, document)


sdt_doc = Document()
sdt_doc.add_paragraph("direct paragraph")
sdt_paragraph = sdt_doc.add_paragraph("inside content control")
sdt = OxmlElement("w:sdt")
sdt_content = OxmlElement("w:sdtContent")
sdt_paragraph._p.getparent().replace(sdt_paragraph._p, sdt)
sdt_content.append(sdt_paragraph._p)
sdt.append(sdt_content)
sdt_doc.save("content-control.docx")
sdt_reopened = Document("content-control.docx")
check("doc.paragraphs omits block content-control text (negative control)",
      "inside content control" not in [paragraph.text for paragraph in sdt_reopened.paragraphs])
walked_text = [paragraph.text for paragraph in iter_body_paragraphs(sdt_reopened.element.body, sdt_reopened)]
check("content-control traversal emits the nested paragraph", "inside content control" in walked_text, walked_text)

# Per-run glyph validation must not let a different referenced font hide a missing glyph.
fixture_cmaps = {"CJK Face": {ord("漢")}, "Latin Face": {ord("A")}}
assigned_runs = [("CJK Face", "漢"), ("Latin Face", "漢")]
pooled_passes = all(any(ord(ch) in cmap for cmap in fixture_cmaps.values())
                    for _, text in assigned_runs for ch in text)
per_run_missing = [(face, ch) for face, text in assigned_runs for ch in text
                   if ord(ch) not in fixture_cmaps[face]]
check("pooled cmap is proven unsafe (negative control)", pooled_passes)
check("per-run cmap check identifies the actual missing glyph",
      per_run_missing == [("Latin Face", "漢")], per_run_missing)

# ---- edit.md guarded cross-run replacement ------------------------------------
SAFE_RUN_CHILDREN = {
    qn("w:rPr"), qn("w:t"), qn("w:tab"), qn("w:br"), qn("w:cr"),
}


def unsafe_run_content(run):
    unsafe = []
    for child in run._r:
        typed_break = child.tag == qn("w:br") and child.get(qn("w:type")) not in (
            None, "textWrapping",
        )
        if child.tag not in SAFE_RUN_CHILDREN or typed_break:
            unsafe.append(child.tag.rsplit("}", 1)[-1])
    return unsafe


def replace_across_runs(paragraph, old, new):
    if not old:
        raise ValueError("old must not be empty")
    runs = list(paragraph.runs)
    text = "".join(run.text for run in runs)
    starts = []
    position = 0
    while (start := text.find(old, position)) != -1:
        starts.append(start)
        position = start + len(old)
    spans = []
    position = 0
    for index, run in enumerate(runs):
        end = position + len(run.text)
        if end > position:
            spans.append((index, position, end))
        position = end
    matches = []
    for start in starts:
        end = start + len(old)
        first, first_start, _ = next(s for s in spans if s[1] <= start < s[2])
        last, last_start, _ = next(s for s in spans if s[1] < end <= s[2])
        matches.append((start, end, first, first_start, last, last_start))
    affected_indexes = {
        index
        for _, _, first, _, last, _ in matches
        for index in range(first, last + 1)
    }
    unsafe = {}
    for index in affected_indexes:
        if children := unsafe_run_content(runs[index]):
            unsafe[index] = children
    if unsafe:
        raise ValueError(f"matched runs contain non-text content: {unsafe}")
    for start, end, first, first_start, last, last_start in reversed(matches):
        prefix = runs[first].text[:start - first_start]
        suffix = runs[last].text[end - last_start:]
        if first == last:
            runs[first].text = prefix + new + suffix
        else:
            runs[first].text = prefix + new
            for index in range(first + 1, last):
                runs[index].text = ""
            runs[last].text = suffix
    return len(starts)


safe_doc = Document()
safe_paragraph = safe_doc.add_paragraph()
safe_first = safe_paragraph.add_run("T")
safe_first.bold = True
safe_paragraph.add_run("B")
safe_paragraph.add_run("D")
check("text-only cross-run match is replaced", replace_across_runs(safe_paragraph, "TBD", "Done") == 1)
check("safe replacement keeps first-run formatting", safe_paragraph.text == "Done" and safe_first.bold)

icon = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 4, 4), False)
icon.clear_with(200)
icon.save("inline-icon.png")
guard_doc = Document()
guard_paragraph = guard_doc.add_paragraph()
guard_run = guard_paragraph.add_run("TBD")
guard_run.add_picture("inline-icon.png")
try:
    replace_across_runs(guard_paragraph, "TBD", "Done")
    rejected_drawing_run = False
except ValueError:
    rejected_drawing_run = True
check("replacement rejects a matched run containing a drawing", rejected_drawing_run)
check(
    "rejected replacement leaves text and drawing untouched",
    guard_run.text == "TBD" and len(guard_run._r.findall(qn("w:drawing"))) == 1,
)


def list_number_num_id(doc):
    """The numId that the ListNumber style binds to in this document part."""
    styles = doc.part.element.body.getparent()  # document.xml root; styles live in another part
    styles_part = doc.part.part_related_by(RT.STYLES)
    for style in styles_part.element.findall(qn("w:style")):
        if style.get(qn("w:styleId")) == "ListNumber":
            numPr = style.find(qn("w:pPr") + "/" + qn("w:numPr"))
            if numPr is not None:
                return int(numPr.find(qn("w:numId")).get(qn("w:val")))
    raise LookupError("ListNumber style has no numPr")


def new_restart_num_id(doc, base_num_id):
    """Clone <w:num> base_num_id with a startOverride so the next list restarts at 1."""
    numbering = doc.part.part_related_by(RT.NUMBERING).element
    source = next(
        n for n in numbering.findall(qn("w:num"))
        if n.get(qn("w:numId")) == str(base_num_id)
    )
    clone = copy.deepcopy(source)
    new_id = max(int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))) + 1
    clone.set(qn("w:numId"), str(new_id))
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    clone.append(override)
    numbering.append(clone)
    return new_id


def numbered_paragraph(doc, text, num_id):
    p = doc.add_paragraph(text, style="List Number")
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    return p


# ---- document A: naive reuse of List Number (second list continues) -----------
doc_a = Document()
doc_a.add_paragraph("First list")
for item in ("one", "two", "three"):
    doc_a.add_paragraph(item, style="List Number")
doc_a.add_paragraph("Second list, same style")
for item in ("four", "five", "six"):
    doc_a.add_paragraph(item, style="List Number")
doc_a.save("continuing.docx")

# ---- document B: second list restarts via cloned numbering definition ---------
doc_b = Document()
doc_b.add_paragraph("First list")
for item in ("one", "two", "three"):
    doc_b.add_paragraph(item, style="List Number")
doc_b.add_paragraph("Second list, restarted")
base_id = list_number_num_id(doc_b)
restart_id = new_restart_num_id(doc_b, base_id)
for item in ("four", "five", "six"):
    numbered_paragraph(doc_b, item, restart_id)
doc_b.save("restarted.docx")

# ---- structural assertions ------------------------------------------------------
numbering_b = doc_b.part.part_related_by(RT.NUMBERING).element
nums_b = numbering_b.findall(qn("w:num"))
check("cloned num entry exists", any(n.get(qn("w:numId")) == str(restart_id) for n in nums_b))
clone_entry = next(n for n in nums_b if n.get(qn("w:numId")) == str(restart_id))
override = clone_entry.find(qn("w:lvlOverride"))
check("clone carries startOverride=1 at ilvl 0",
      override is not None and override.get(qn("w:ilvl")) == "0"
      and override.find(qn("w:startOverride")).get(qn("w:val")) == "1")
second_list_num_ids = [
    p._p.find(qn("w:pPr") + "/" + qn("w:numPr") + "/" + qn("w:numId")).get(qn("w:val"))
    for p in doc_b.paragraphs[-3:]
]
check("restarted paragraphs reference the cloned numId", set(second_list_num_ids) == {str(restart_id)}, second_list_num_ids)

# ---- rendered proof via LibreOffice --------------------------------------------
subprocess.run(
    ["soffice", "--headless", "--convert-to", "pdf", "--outdir", ".", "continuing.docx", "restarted.docx"],
    check=True, capture_output=True, timeout=180,
)


def second_list_numbers(pdf_path):
    text = " ".join(page.get_text() for page in fitz.open(pdf_path))
    # take the rendered numbers in front of the second list's item words
    return [text.split(word)[0].split()[-1] for word in ("four", "five", "six")]


cont = second_list_numbers("continuing.pdf")
restart = second_list_numbers("restarted.pdf")
print("continuing.docx renders second list as:", cont)
print("restarted.docx renders second list as:", restart)
check("plain style reuse continues the sequence (negative control)", cont == ["4.", "5.", "6."], cont)
check("cloned definition restarts the second list at 1", restart == ["1.", "2.", "3."], restart)

# ---- scenes.md snippet: keep_table_together for the signature block -------------
doc_c = Document()
doc_c.add_heading("Contract", level=1)
# fill most of page 1 so a tall signature table would otherwise straddle the page break
for _ in range(24):
    doc_c.add_paragraph("Filler paragraph to push the signature block toward the page break. " * 3)
sig = doc_c.add_table(rows=4, cols=2)
labels = [("甲方（盖章）", "乙方（盖章）"), ("签字", "签字"), ("日期", "日期"), ("备注", "备注")]
for r, pair in enumerate(labels):
    sig.cell(r, 0).text, sig.cell(r, 1).text = pair


def keep_table_together(table):
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))   # a row never splits mid-row
    for row in table.rows[:-1]:
        for cell in row.cells:
            for par in cell.paragraphs:
                par.paragraph_format.keep_with_next = True  # row sticks to the next row


keep_table_together(sig)
doc_c.save("signature.docx")
reopened = Document("signature.docx")
check("every signature row carries cantSplit",
      all(row._tr.find(qn("w:trPr")) is not None and row._tr.find(qn("w:trPr")).find(qn("w:cantSplit")) is not None
          for row in reopened.tables[0].rows))
subprocess.run(
    ["soffice", "--headless", "--convert-to", "pdf", "--outdir", ".", "signature.docx"],
    check=True, capture_output=True, timeout=180,
)
pages_with_labels = [
    page.number for page in fitz.open("signature.pdf")
    if "甲方（盖章）" in page.get_text() and "备注" in page.get_text()
]
check("rendered signature table stays on one page", len(pages_with_labels) == 1, pages_with_labels)

print("\n" + ("ALL DOCX FIXTURES PASSED" if not failures else f"{len(failures)} FAILURES: {failures}"))
sys.exit(0 if not failures else 1)
